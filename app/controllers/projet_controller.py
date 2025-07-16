from flask import Blueprint, render_template, jsonify, request, send_file
from ..models.projet import Projet
from app import db
from docx import Document
from io import BytesIO
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import pandas as pd
from app.utils.token_required import token_required
from app.utils.role_required import admin_required

bp = Blueprint('projets', __name__)



EXTENSIONS_AUTORISEE = ["xls", "xlsx"]

@bp.route('/projets', methods=['POST'])
@token_required
def create_projet(current_user):
    data = request.get_json()
    date_debut = data.get('date_debut')
    if not date_debut:
        date_debut = datetime.now()
        
    date_fin = data.get('date_fin')
    if not date_fin:
        date_fin = datetime.now()
    projet = Projet(
        nom_client=data.get('nom_client'),
        nom_projet=data.get('nom_projet'),
        ville=data.get('ville'),
        pays=data.get('pays'),
        adresse_client=data.get('adresse_client'),
        contacts_client=data.get('contacts_client'),
        domaine_expertise=data.get('domaine_expertise'),
        metier=data.get('metier'),
        desc_courte=data.get('desc_courte'),
        desc_longue=data.get('desc_longue'),
        resultat_impact=data.get('resultat_impact'),
        contact_ressource=data.get('contact_ressource'),
        equipe_projet=data.get('equipe_projet'),
        nb_employes_mission=safe_int(data.get('nb_employes_mission')),
        consultants_associes=data.get('consultants_associes'),
        nb_employes_consultants=safe_int(data.get('nb_employes_consultants')),
        cadres_societe=data.get('cadres_societe'),
        date_debut= date_debut,
        date_fin=date_fin,
        contient_documents=data.get('contient_documents'),
        cout_projet=data.get('cout_projet'),
        projet_confidentiel=data.get('projet_confidentiel'),
        client_confidentiel=data.get('client_confidentiel'),
        desc_anglaise=data.get('desc_anglaise'),
        sous_domaines=data.get('sous_domaines')
    )
    db.session.add(projet)
    db.session.commit()
    return jsonify(projet.to_dict()), 201



@bp.route('/projets', methods=['GET'])
@token_required
def get_projets(current_user):
    projets = Projet.query.all()
    return jsonify([p.to_dict() for p in projets]), 200



@bp.route('/projets/<int:id>', methods=['GET'])
@token_required
def get_projet(current_user, id):
    projet = Projet.query.get_or_404(id)
    return jsonify(projet.to_dict()), 200


@bp.route('/projets/<int:id>', methods=['PUT'])
@admin_required
def update_projet(current_user, id):
    data = request.get_json()
    projet = Projet.query.get_or_404(id)

    for key, value in data.items():
        if hasattr(projet, key):
            setattr(projet, key, value)

    db.session.commit()
    return jsonify(projet.to_dict()), 200


@bp.route('/projets/<int:id>', methods=['DELETE'])
@admin_required
def delete_projet(current_user, id):
    projet = Projet.query.get_or_404(id)
    db.session.delete(projet)
    db.session.commit()
    return jsonify({'message': 'Projet supprimé'}), 200


@bp.route('/delete-projets-selected', methods=["DELETE"])
@token_required
def deleteProjectsSelected(current_user):
    """
    Télécharge plusieurs projets sélectionnés au format FS
    Chaque projet est sur une page différente
    """
    # Récupérer les IDs des projets depuis les paramètres de requête
    project_ids = []
    
    # Gérer le format ids[] envoyé par axios
    ids_from_params = request.args.getlist('ids[]')
    if ids_from_params:
        for id_str in ids_from_params:
            try:
                project_ids.append(int(id_str))
            except ValueError:
                continue
    
    # Si pas d'IDs dans les paramètres, essayer de les récupérer depuis le JSON body
    if not project_ids and request.is_json:
        data = request.get_json()
        project_ids = data.get('project_ids', [])
    
    if not project_ids:
        return jsonify({'error': 'Aucun projet sélectionné'}), 400
    
    for index, project_id in enumerate(project_ids):
        try:
            # Récupérer le projet
            projet = Projet.query.get_or_404(project_id)
            
            if not projet:
                continue
            db.session.delete(projet)
            db.session.commit()
                
        except Exception as e:
            print(f"Erreur lors du traitement du projet {project_id}: {str(e)}")
            continue
    
    return jsonify({'message': 'Projet supprimé'}), 200

@bp.route('/projets/stats', methods=['GET'])
@token_required
def stats(current_user):
    projets = Projet.query.all()

    stats = {
        "total": 0,
        "terminés": 0,
        "en_cours": 0,
        "à_venir": 0
    }

    for projet in projets:
        statut = formatage_date(projet.date_debut, projet.date_fin)
        stats["total"] += 1
        if statut == "projet terminé":
            stats["terminés"] += 1
        elif statut == "projet en cours":
            stats["en_cours"] += 1
        elif statut == "projet à venir":
            stats["à_venir"] += 1



@bp.route('/download-word-bm/<int:id>')
@token_required
def download_one_bm_word(current_user, id):
    doc = Document()
    # Exemple de données
    projet = Projet.query.get_or_404(id)
    
    print(projet)
   


    table = doc.add_table(rows=8, cols=2, style='Table Grid')
    # Récupère les cellules de la première ligne
    hdr_cells = table.rows[0].cells

    # Fusionne les 2 cellules en une seule (de la cellule 0 à la cellule 2)
    merged_cell = hdr_cells[0].merge(hdr_cells[1])
    
    # Mets du texte dans la cellule fusionnée
    merged_cell.text = f""
        

    # Centrage vertical
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    
    paragraph = merged_cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(12)
    run.bold = True

    # Appliquer un fond orange (code hex sans le #)
    set_cell_background(merged_cell, "3d8547")  # Vert Primary
    


    # LIGNE 1
    row_1 = table.rows[1].cells
    # Cellule gauche
    p1_0 = row_1[0].paragraphs[0]
    p1_0.add_run("Nom de la mission : ").bold = True
    p1_0.add_run(f"{projet.nom_projet}")
    # Cellule droite
    p1_1 = row_1[1].paragraphs[0]
    p1_1.add_run("Valeur du contrat (en francs CFA) : ").bold = True
    p1_1.add_run(f"{format_cout_projet(projet.cout_projet)}")

    # LIGNE 2
    row_2 = table.rows[2].cells
    # Cellule gauche
    p2_0 = row_2[0].paragraphs[0]
    p2_0.add_run("Pays : ").bold = True
    p2_0.add_run(f"{projet.pays}\n")
    p2_0.add_run("Lieu : ").bold = True
    p2_0.add_run(f"{projet.ville}")
    # Cellule droite
    p2_1 = row_2[1].paragraphs[0]
    p2_1.add_run("Durée de la mission (mois) : ").bold = True
    p2_1.add_run(f"{calcul_duree_mission_en_mois(projet.date_debut, projet.date_fin)} ")

    # LIGNE 3
    row_3 = table.rows[3].cells
    # Cellule gauche
    p3_0 = row_3[0].paragraphs[0]
    p3_0.add_run("Nom du client : ").bold = True
    p3_0.add_run(f"{projet.nom_client}")
    # Cellule droite
    p3_1 = row_3[1].paragraphs[0]
    p3_1.add_run("Nombre total d'employés/mois ayant participé à la mission : ").bold = True
    p3_1.add_run(f"{projet.nb_employes_mission}")

    # LIGNE 4
    row_4 = table.rows[4].cells
    # Cellule gauche
    p4_0 = row_4[0].paragraphs[0]
    p4_0.add_run("Adresse : ").bold = True
    p4_0.add_run(f"{projet.adresse_client}")
    # Cellule droite fusionnée (ligne 4-5 col 1)
    merge_cell_row_4 = row_4[1].merge(table.cell(5, 1))
    p4_1 = merge_cell_row_4.paragraphs[0]
    p4_1.add_run("Noms des cadres professionnels de votre société employés et fonctions exécutées : ").bold = True
    p4_1.add_run(f"{projet.cadres_societe}")

    # LIGNE 5
    row_5 = table.rows[5].cells
    p5_0 = row_5[0].paragraphs[0]
    p5_0.add_run("Date de démarrage (mois/année) : ").bold = True
    p5_0.add_run(f"{format_date_mois_annee(projet.date_debut)}\n")
    p5_0.add_run("Date d’achèvement (mois/année) : ").bold = True
    p5_0.add_run(f"{format_date_mois_annee(projet.date_fin)}")
    # Ne rien écrire dans row_5[1], car elle est fusionnée

    # LIGNE 6
    row_6 = table.rows[6].cells
    merged_cell_6 = row_6[0].merge(row_6[1])
    p6 = merged_cell_6.paragraphs[0]
    p6.add_run("Description du projet : ").bold = True
    p6.add_run(f"{projet.desc_courte}")

    # LIGNE 7
    row_7 = table.rows[7].cells
    merged_cell_7 = row_7[0].merge(row_7[1])
    p7 = merged_cell_7.paragraphs[0]
    p7.add_run("Description des services effectivement rendus par votre personnel dans le cadre de la mission : ").bold = True
    p7.add_run(f"{projet.desc_longue}")



    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"projet_{projet.nom_projet}_format_BM.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@bp.route('/download-bm-projets-selected', methods=["GET", "POST"])
@token_required
def downloadBmProjectsSelected(current_user):
    """
    Télécharge plusieurs projets sélectionnés au format Word BM
    Chaque projet est sur une page différente
    """
    # Récupérer les IDs des projets depuis les paramètres de requête
    # Flask ne reconnaît pas automatiquement le format ids[] d'axios
    project_ids = []
    
    # Gérer le format ids[] envoyé par axios
    # request.args.getlist('ids[]') récupère toutes les valeurs avec la clé 'ids[]'
    ids_from_params = request.args.getlist('ids[]')
    if ids_from_params:
        for id_str in ids_from_params:
            try:
                project_ids.append(int(id_str))
            except ValueError:
                continue
    
    # Si pas d'IDs dans les paramètres, essayer de les récupérer depuis le JSON body
    if not project_ids and request.is_json:
        data = request.get_json()
        project_ids = data.get('project_ids', [])
    
    if not project_ids:
        return jsonify({'error': 'Aucun projet sélectionné'}), 400
    
    # Créer un nouveau document
    doc = Document()
    
    # Pour chaque projet sélectionné
    for index, project_id in enumerate(project_ids):
        try:
            # Récupérer le projet
            projet = Projet.query.get(project_id)
            if not projet:
                continue
                
            # Ajouter un saut de page entre les projets (sauf pour le premier)
            if index > 0:
                doc.add_page_break()
            
            # Créer le tableau pour ce projet
            table = doc.add_table(rows=8, cols=2, style='Table Grid')
            
            # LIGNE 0 - En-tête fusionné
            hdr_cells = table.rows[0].cells
            merged_cell = hdr_cells[0].merge(hdr_cells[1])
            merged_cell.text = f""
            
            # Centrage vertical et horizontal
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = merged_cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            run.font.size = Pt(12)
            run.bold = True
            
            # Appliquer un fond orange
            set_cell_background(merged_cell, "3d8547")
            
         

            # LIGNE 1
            row_1 = table.rows[1].cells
            # Cellule gauche
            p1_0 = row_1[0].paragraphs[0]
            p1_0.add_run("Nom de la mission : ").bold = True
            p1_0.add_run(f"{projet.nom_projet}")
            # Cellule droite
            p1_1 = row_1[1].paragraphs[0]
            p1_1.add_run("Valeur du contrat (en francs CFA) : ").bold = True
            p1_1.add_run(f"{format_cout_projet(projet.cout_projet)}")

            # LIGNE 2
            row_2 = table.rows[2].cells
            # Cellule gauche
            p2_0 = row_2[0].paragraphs[0]
            p2_0.add_run("Pays : ").bold = True
            p2_0.add_run(f"{projet.pays}\n")
            p2_0.add_run("Lieu : ").bold = True
            p2_0.add_run(f"{projet.ville}")
            # Cellule droite
            p2_1 = row_2[1].paragraphs[0]
            p2_1.add_run("Durée de la mission (mois) : ").bold = True
            p2_1.add_run(f"{calcul_duree_mission_en_mois(projet.date_debut, projet.date_fin)}")

            # LIGNE 3
            row_3 = table.rows[3].cells
            # Cellule gauche
            p3_0 = row_3[0].paragraphs[0]
            p3_0.add_run("Nom du client : ").bold = True
            p3_0.add_run(f"{projet.nom_client}")
            # Cellule droite
            p3_1 = row_3[1].paragraphs[0]
            p3_1.add_run("Nombre total d'employés/mois ayant participé à la mission : ").bold = True
            p3_1.add_run(f"{projet.nb_employes_mission}")

            # LIGNE 4
            row_4 = table.rows[4].cells
            # Cellule gauche
            p4_0 = row_4[0].paragraphs[0]
            p4_0.add_run("Adresse : ").bold = True
            p4_0.add_run(f"{projet.adresse_client}")
            # Cellule droite fusionnée (ligne 4-5 col 1)
            merge_cell_row_4 = row_4[1].merge(table.cell(5, 1))
            p4_1 = merge_cell_row_4.paragraphs[0]
            p4_1.add_run("Noms des cadres professionnels de votre société employés et fonctions exécutées : ").bold = True
            p4_1.add_run(f"{projet.cadres_societe}")

            # LIGNE 5
            row_5 = table.rows[5].cells
            p5_0 = row_5[0].paragraphs[0]
            p5_0.add_run("Date de démarrage (mois/année) : ").bold = True
            p5_0.add_run(f"{format_date_mois_annee(projet.date_debut)}\n")
            p5_0.add_run("Date d’achèvement (mois/année) : ").bold = True
            p5_0.add_run(f"{format_date_mois_annee(projet.date_fin)}")
            # Ne rien écrire dans row_5[1], car elle est fusionnée

            # LIGNE 6
            row_6 = table.rows[6].cells
            merged_cell_6 = row_6[0].merge(row_6[1])
            p6 = merged_cell_6.paragraphs[0]
            p6.add_run("Description du projet : ").bold = True
            p6.add_run(f"{projet.desc_courte}")

            # LIGNE 7
            row_7 = table.rows[7].cells
            merged_cell_7 = row_7[0].merge(row_7[1])
            p7 = merged_cell_7.paragraphs[0]
            p7.add_run("Description des services effectivement rendus par votre personnel dans le cadre de la mission : ").bold = True
            p7.add_run(f"{projet.desc_longue}")


        except Exception as e:
            print(f"Erreur lors du traitement du projet {project_id}: {str(e)}")
            continue
    
    # Sauvegarder le document dans un flux de bytes
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Retourner le fichier
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"projets_selection_{len(project_ids)}_format_BM.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )




@bp.route('/download-word-fs/<int:id>')
@token_required
def download_one_fs_word(current_user,id):
    doc = Document()
    # Exemple de données
    projet = Projet.query.get_or_404(id)
    
    print(projet)
   


    table = doc.add_table(rows=2, cols=5, style='Table Grid')

    #============ remaniement la ligne corresponds à l'index 0 (seulement pour les Format de synthèse)===============
    # LIGNE 1
    row_1 = table.rows[0].cells
    headers = ["Date", "Nom du client", "Nom de la mission", "Brève description", "Valeur du contrat"]

    for i, text in enumerate(headers):
        cell = row_1[i]
        set_cell_background(cell, "3d8547")  # vert
        paragraph = cell.paragraphs[0]
        paragraph.clear()  # Nettoyer le paragraphe
        run = paragraph.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc

    
     # LIGNE 2
    row_2 = table.rows[1].cells
    row_2[0].text = f"{format_date_mois_annee(projet.date_debut) } - {format_date_mois_annee(projet.date_fin)}"
    row_2[1].text = f"{ projet.nom_client }"
    row_2[2].text = f"{ projet.nom_projet }"
    row_2[3].text = f"{ projet.desc_courte }"
    row_2[4].text = f"{ format_cout_projet(projet.cout_projet) }"
    
    
    

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"projet_{projet.nom_projet}_format_FS.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@bp.route('/download-fs-projets-selected', methods=["GET", "POST"])
@token_required
def downloadFsProjectsSelected(current_user):
    """
    Télécharge plusieurs projets sélectionnés au format FS
    Chaque projet est sur une page différente
    """
    # Récupérer les IDs des projets depuis les paramètres de requête
    project_ids = []
    
    # Gérer le format ids[] envoyé par axios
    ids_from_params = request.args.getlist('ids[]')
    if ids_from_params:
        for id_str in ids_from_params:
            try:
                project_ids.append(int(id_str))
            except ValueError:
                continue
    
    # Si pas d'IDs dans les paramètres, essayer de les récupérer depuis le JSON body
    if not project_ids and request.is_json:
        data = request.get_json()
        project_ids = data.get('project_ids', [])
    
    if not project_ids:
        return jsonify({'error': 'Aucun projet sélectionné'}), 400
    
    # Créer un nouveau document
    doc = Document()
    
    # Pour chaque projet sélectionné
    for index, project_id in enumerate(project_ids):
        try:
            # Récupérer le projet
            projet = Projet.query.get(project_id)
            if not projet:
                continue
                
            # Ajouter un saut de page entre les projets (sauf pour le premier)
            if index > 0:
                doc.add_page_break()
            
            # Créer le tableau pour ce projet (3 lignes, 5 colonnes)
                    
            table = doc.add_table(rows=2, cols=5, style='Table Grid')

            #============ remaniement la ligne corresponds à l'index 0 (seulement pour les Format de synthèse)===============
            # LIGNE 1
            row_1 = table.rows[0].cells
            headers = ["Date", "Nom du client", "Nom de la mission", "Brève description", "Valeur du contrat"]

            for i, text in enumerate(headers):
                cell = row_1[i]
                set_cell_background(cell, "3d8547")  # vert
                paragraph = cell.paragraphs[0]
                paragraph.clear()  # Nettoyer le paragraphe
                run = paragraph.add_run(text)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc

            
            # LIGNE 2
            row_2 = table.rows[1].cells
            row_2[0].text = f"{format_date_mois_annee(projet.date_debut) } - {format_date_mois_annee(projet.date_fin)}"
            row_2[1].text = f"{ projet.nom_client }"
            row_2[2].text = f"{ projet.nom_projet }"
            row_2[3].text = f"{ projet.desc_courte }"
            row_2[4].text = f"{ format_cout_projet(projet.cout_projet) }"
            
        except Exception as e:
            print(f"Erreur lors du traitement du projet {project_id}: {str(e)}")
            continue
    
    # Sauvegarder le document dans un flux de bytes
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Retourner le fichier
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"projets_selection_{len(project_ids)}_format_FS.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@bp.route('/download-word-fs')
@token_required
def download_all_fs_word(current_user):
    doc = Document()
    # Exemple de données
    projets = Projet.query.all()
    

    nb_ligne_a_ajouter = 1
    nb_rows_table = len(projets)  + nb_ligne_a_ajouter

    table = doc.add_table(rows=nb_rows_table, cols=5, style='Table Grid')
    
    row_1 = table.rows[0].cells

    headers = ["Date", "Nom du client", "Nom de la mission", "Brève description", "Valeur du contrat"]

    for i, text in enumerate(headers):
        cell = row_1[i]
        set_cell_background(cell, "3d8547")  # vert
        paragraph = cell.paragraphs[0]
        paragraph.clear()  # Nettoyer le paragraphe
        run = paragraph.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc
        
    for index, projet in enumerate(projets, start=nb_ligne_a_ajouter):
        row_table = table.rows[index].cells
        row_table[0].text = f"{format_date_mois_annee(projet.date_debut) } - {format_date_mois_annee(projet.date_fin)}"
        row_table[1].text = f"{ projet.nom_client }"
        row_table[2].text = f"{ projet.nom_projet }"
        row_table[3].text = f"{ format_cout_projet(projet.cout_projet) }"
        
        
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)


    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"liste_projets_format_FS.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )



@bp.get('/download-excel')
@token_required
def download_all_excel_old_format(current_user):
    projets = Projet.query.all()
   # Initialisation des listes
    num_projet = []
    nom_client = []
    nom_projet = []
    pays = []
    ville = []
    adresse_postale = []
    contacts = []
    domaine_expertise = []
    metier = []
    desc_une_phrase = []
    desc_paragraphe = []
    resultat_impact = []
    contact_ressource = []
    equipe_projet = []
    nb_employes = []
    consultants_associes = []
    nb_employes_consultants = []
    cadres_societe = []
    date_demarrage = []
    date_achevement = []
    contient_documents = []
    cout_projet = []
    projet_confidentiel = []
    client_confidentiel = []
    desc_anglais = []
    sous_domaines = []

    # Remplissage
    for projet in projets:
        num_projet.append(projet.id)
        nom_client.append(projet.nom_client)
        nom_projet.append(projet.nom_projet)
        pays.append(projet.pays or "")
        ville.append(projet.ville or "")
        adresse_postale.append(projet.adresse_client)
        contacts.append(projet.contacts_client)
        domaine_expertise.append(projet.domaine_expertise)
        metier.append(projet.metier)
        desc_une_phrase.append(projet.desc_courte)
        desc_paragraphe.append(projet.desc_longue)
        resultat_impact.append(projet.resultat_impact)
        contact_ressource.append(projet.contact_ressource)
        equipe_projet.append(projet.equipe_projet)
        nb_employes.append(projet.nb_employes_mission)
        consultants_associes.append(projet.consultants_associes)
        nb_employes_consultants.append(projet.nb_employes_consultants)
        cadres_societe.append(projet.cadres_societe)
        date_demarrage.append(projet.date_debut.strftime("%d/%m/%Y") if projet.date_debut else "")
        date_achevement.append(projet.date_fin.strftime("%d/%m/%Y") if projet.date_fin else "")
        contient_documents.append("Oui" if projet.contient_documents else "Non")
        cout_projet.append(projet.cout_projet)
        projet_confidentiel.append("Oui" if projet.projet_confidentiel else "Non")
        client_confidentiel.append("Oui" if projet.client_confidentiel else "Non")
        desc_anglais.append(projet.desc_anglaise)
        sous_domaines.append(projet.sous_domaines)

    # Construction du dictionnaire final
    data = {
        'Numéro du projet': num_projet,
        'Nom du client': nom_client,
        'Nom du projet': nom_projet,
        'Pays': pays,
        'Ville': ville,
        'Adresse postale et géographique du client': adresse_postale,
        'Contacts téléphoniques et adresse électronique du client': contacts,
        'Domaine d\'expertise': domaine_expertise,
        'Métier': metier,
        'Description du projet en une phrase': desc_une_phrase,
        'Description du projet en un paragraphe': desc_paragraphe,
        'Résultat et impact obtenus par le projet': resultat_impact,
        'Contact Ressource Athari': contact_ressource,
        'Equipe projet': equipe_projet,
        'Nombre total d’employés / mois ayant participé à la Mission ': nb_employes,
        'Noms des consultants associés/partenaires éventuels': consultants_associes,
        'Nombre d\'employés/mois fournis par les consultants associés :': nb_employes_consultants,
        'Noms des cadres professionnels de votre société employés et fonctions :': cadres_societe,
        'Date de démarrage': date_demarrage,
        'Date d\'achèvement': date_achevement,
        'Contient des documents': contient_documents,
        'Coût du projet': cout_projet,
        'Projet confidentiel': projet_confidentiel,
        'Client confidentiel': client_confidentiel,
        'Description du projet en Anglais': desc_anglais,
        'Sous-domaines du projet': sous_domaines
    }

    df = pd.DataFrame(data)
    

    
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Base de donnée')
        
            # Récupérer le classeur et la feuille
        workbook = writer.book
        worksheet = writer.sheets['Base de donnée']

        # Définir le style d'en-tête orange
        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'fg_color': '#3d8547',  # orange
            'border': 1
        })
        
        
        # Appliquer le style à chaque cellule de l'en-tête
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
        
    output.seek(0)  # Revenir au début du fichier en mémoire

    # Envoyer le fichier au client
    return send_file(
        output,
        as_attachment=True,
        download_name='Base de donnée Athari.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    


@bp.route('/upload-file', methods=["POST"])
@token_required
def uploadFile(current_user):
    # Validation des fichiers
    if 'excel_file' not in request.files:
        return "Aucun fichier équivalent trouvé dans le champ excel_file"

    excelFile = request.files["excel_file"]

    if excelFile.filename == "":
        return "Aucun fichier selectionné", 400
    
    try:
        # Lecture du fichier Excel en mémoire
        data = pd.read_excel(BytesIO(excelFile.read()))
        print(data.dtypes)
        print(data.columns.tolist())
        for index, row in data.iterrows():

            # row est une Series => il faut utiliser row.iloc[i] pour accéder par position
            projet = Projet(
            nom_client=safe_str(row["Nom du client"]),
            nom_projet=safe_str(row["Nom du projet"]),
            pays=safe_str(row["Pays"]),
            ville=safe_str(row["Ville"]),
            adresse_client=safe_str(row["Adresse postale et géographique du client"]),
            contacts_client=safe_str(row["Contacts téléphoniques et adresse électronique du client"]),
            domaine_expertise=safe_str(row["Domaine d'expertise"]),
            metier=safe_str(row["Métier"]),
            desc_courte=safe_str(row["Description du projet en une phrase"]),
            desc_longue=safe_str(row["Description du projet en un paragraphe"]),
            resultat_impact=safe_str(row["Résultat et impact obtenus par le projet"]),
            contact_ressource=safe_str(row["Contact Ressource Athari"]),
            equipe_projet=safe_str(row["Equipe projet"]),
            nb_employes_mission=safe_int(row["Nombre total d’employés / mois ayant participé à la Mission "]),
            consultants_associes=safe_str(row["Noms des consultants associés/partenaires éventuels"]),
            nb_employes_consultants=safe_int(row["Nombre d'employés/mois fournis par les consultants associés :"]),
            cadres_societe=safe_str(row["Noms des cadres professionnels de votre société employés et fonctions :"]),
            date_debut=safe_date(row["Date de démarrage"]),
            date_fin=safe_date(row["Date d'achèvement"]),
            contient_documents=safe_bool(row["Contient des documents"]),
            cout_projet=safe_str(row["Coût du projet"]),
            projet_confidentiel=safe_bool(row["Projet confidentiel"]),
            client_confidentiel=safe_bool(row["Client confidentiel"]),
            desc_anglaise=safe_str(row["Description du projet en Anglais"]),
            sous_domaines=safe_str(row["Sous-domaines du projet"])
)
            db.session.add(projet)
            
            
        db.session.commit()
        return "Fichier traité avec succès", 200

    except Exception as e:
        import traceback
        print(traceback.format_exc())  # Affiche toute la trace d'erreur
        return f"Erreur lors de la lecture du fichier : {str(e)}", 500



    
    
    
    
    

    
    
    

def calcul_duree_mission_en_mois(date_debut, date_fin):
    start = datetime.strptime(str(date_debut), '%Y-%m-%d')
    end = datetime.strptime(str(date_fin), '%Y-%m-%d')
    
    mois = (end.year - start.year) * 12 + (end.month - start.month)

    # Si le jour de fin est >= au jour de début, on ajoute 1 mois
    if end.day >= start.day:
        mois += 1
    
    return f"{max(0, mois)} mois"  # on évite les durées négatives
 

def format_cout_projet(cout_projet):
    try:
        montant = float(cout_projet)
        montant = int(montant)
        return f"{montant:,} FCFA".replace(",", " ")  # format avec espace tous les milliers
    except (ValueError, TypeError):
        return str(cout_projet)

def format_date_mois_annee(date_val):
    try:
        return date_val.strftime("%m/%Y")
    except AttributeError:
        return str(date_val)
   

def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)  # couleur hex sans #
    tcPr.append(shd)
    

def autoriser_typ_fichiers(filename):
    return '.' in filename and filename.rsplit('.', 1)[1] in EXTENSIONS_AUTORISEE



def formatage_date(date_debut, date_fin):
    aujourd_hui = datetime.today().date()
    debut = datetime.strptime(date_debut, "%Y-%m-%d").date()
    fin = datetime.strptime(date_fin, "%Y-%m-%d").date()

    if fin < aujourd_hui:
        return "projet terminé"
    elif debut > aujourd_hui:
        return "projet à venir"
    else:
        return "projet en cours"


def safe_int(value):
    try:
        if pd.isna(value):
            return None
        return int(float(value))  # permet aussi de convertir "3.0" en 3
    except (ValueError, TypeError):
        return None


def safe_str(value):
    return str(value).strip() if pd.notna(value) else None

def safe_bool(value):
    if pd.isna(value):
        return False
    if isinstance(value, str):
        return value.strip().lower() in ['true', '1', 'oui', 'yes']
    return bool(value)

def safe_date(value):
    if isinstance(value, datetime):
        return value
    try:
        return datetime.strptime(value, "%d/%m/%Y")
    except (ValueError, TypeError):
        return None  