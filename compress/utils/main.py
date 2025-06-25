import io
import os
import zipfile
from datetime import datetime

from docx import Document
from flask import Flask, render_template, request, url_for, redirect, send_file
import pandas as pd



PATH_FILE = "static/uploads"
OUTPUT_FOLDER = "output"

EXTENSIONS_AUTORISEE = ["xls", "xlsx"]
# ------------------------ UTILS -------------------

def excelToWorld(filePath, date_heure_str):
    data = pd.read_excel(filePath)

    # Création d'un dossier unique pour ahy les outputs
    newFolderForStockDoc = f"{OUTPUT_FOLDER}/{date_heure_str}"
    os.makedirs(newFolderForStockDoc)

    print("------")
    for index, row in data.iterrows():
        print(row)
        doc = Document()

        doc.add_paragraph(f"Ligne {index + 1}", style="Heading 2")

        table = doc.add_table(rows=2, cols=len(row), style='Table Grid')
        print(row.index)
        hdr_cells = table.rows[0].cells
        line2 = table.rows[1].cells
        for i, col_name in enumerate(row.index):
            hdr_cells[i].text = f"{col_name}"
            line2[i].text = f"{row[col_name]}"

        doc.save(f"{newFolderForStockDoc}/ligne_{index + 1}.docx")


def autoriser_typ_fichiers(filename):
    return '.' in filename and filename.rsplit('.', 1)[1] in EXTENSIONS_AUTORISEE


def get_directory_structure(rootdir):
    structure = []
    for item in os.listdir(rootdir):
        path = os.path.join(rootdir, item)
        if os.path.isdir(path):
            structure.append({
                'name': item,
                'type': 'folder',
                'children': get_directory_structure(path)
            })
        else:
            structure.append({
                'name': item,
                'type': 'file'
            })
    return structure



def creer_zip_en_memoire(dossier_source):
    mem_zip = io.BytesIO()

    with zipfile.ZipFile(mem_zip, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for dossier_racine, _, fichiers in os.walk(dossier_source):
            for fichier in fichiers:
                chemin_complet = os.path.join(dossier_racine, fichier)
                chemin_relatif = os.path.relpath(chemin_complet, dossier_source)
                zipf.write(chemin_complet, chemin_relatif)

    mem_zip.seek(0)  # Revenir au début du fichier en mémoire
    return mem_zip


